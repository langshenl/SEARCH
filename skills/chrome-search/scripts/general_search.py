#!/usr/bin/env python3
"""
通用网页搜索工具：对任意网站执行关键词搜索，抓取第1页结果
模式1（默认）：生成Word文档
模式2（--analyze）：流式输出+每篇摘要，AI直接分析

用法:
  python3 general_search.py <搜索URL> <关键词> [输出目录] [链接选择器] [正文选择器] [analyze]
"""
import subprocess
import json
import time
import sys
import os
import re
import urllib.parse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CDP = "http://localhost:3456"
CONCURRENCY = 5
MAX_WAIT_READY = 3
DEFAULT_MAX_RETRIES = 2

DEFAULT_LINK_SELECTORS = '.news-list .item-title a, .search-result a'
DEFAULT_BODY_SELECTORS = 'TRS_Editor|bt-article|article-content|content|article|main|body'

# 全局记录已打开的tab，用于异常时清理
_open_tabs = []


def normalize_url(url):
    """
    归一化URL，去除百度跳转参数等，保留真实目标URL
    用于去重和调试
    """
    if not url:
        return ''

    # 百度跳转链接解码
    if 'baidu.com/link?url=' in url:
        try:
            encoded = url.split('baidu.com/link?url=')[1].split('&')[0]
            decoded = urllib.parse.unquote(encoded)
            if decoded.startswith('http'):
                return decoded
        except Exception:
            pass

    # 去除UTM参数
    url = re.sub(r'[?&](utm_|fbclid|gclid|ref|cid|source|from)=[^&]*', '', url)
    # 去除尾部空格
    url = url.strip().rstrip('/')
    return url


def dedup_links(links):
    """
    基于归一化URL去重，保留首次出现的链接
    """
    seen = set()
    deduped = []
    for link in links:
        url = link['href']
        norm = normalize_url(url)
        if norm and norm not in seen:
            seen.add(norm)
            link['_norm_url'] = norm
            deduped.append(link)
        elif not norm:
            # 无法归一化的URL直接保留
            seen.add(url)
            deduped.append(link)
    return deduped


def cdp_new(url):
    """创建新tab，返回targetId"""
    global _open_tabs
    try:
        resp = subprocess.run(
            ['curl', '-s', f'{CDP}/new?url={url}'],
            capture_output=True, text=True, timeout=25
        )
        result = json.loads(resp.stdout.strip())
        target_id = result.get('targetId', '')
        if target_id:
            _open_tabs.append(target_id)
        return target_id
    except Exception as e:
        print(f"    ⚠️ cdp_new失败: {e}", flush=True)
        return ''


def cdp_info(target_id):
    """获取tab信息"""
    try:
        resp = subprocess.run(
            ['curl', '-s', f'{CDP}/info?target={target_id}'],
            capture_output=True, text=True, timeout=10
        )
        return json.loads(resp.stdout.strip())
    except Exception as e:
        print(f"    ⚠️ cdp_info失败: {e}", flush=True)
        return {}


def cdp_eval(target_id, script, timeout=35):
    """执行JS，返回结果字符串"""
    try:
        resp = subprocess.run(
            ['curl', '-s', '-X', 'POST', f'{CDP}/eval?target={target_id}', '-d', script],
            capture_output=True, text=True, timeout=timeout
        )
        obj = json.loads(resp.stdout.strip())
        return obj.get('value', '')
    except json.JSONDecodeError as e:
        print(f"    ⚠️ cdp_eval JSON解析失败: {e}", flush=True)
        return ''
    except Exception as e:
        print(f"    ⚠️ cdp_eval执行失败: {e}", flush=True)
        return ''


def cdp_close(target_id):
    """关闭tab（快速失败，不阻塞）"""
    global _open_tabs
    try:
        # 使用短timeout，避免在清理时卡住
        subprocess.run(['curl', '-s', f'{CDP}/close?target={target_id}'], timeout=3)
        if target_id in _open_tabs:
            _open_tabs.remove(target_id)
    except Exception:
        # 清理失败时直接移除记录，避免残留
        if target_id in _open_tabs:
            _open_tabs.remove(target_id)


def cleanup_all_tabs():
    """强制清理所有已打开的tab（异常保护，快速失败）"""
    global _open_tabs
    if _open_tabs:
        tabs_to_close = _open_tabs[:]
        _open_tabs.clear()
        # fire-and-forget 方式关闭所有tab
        for tid in tabs_to_close:
            try:
                subprocess.run(
                    ['curl', '-s', f'{CDP}/close?target={tid}'],
                    timeout=2
                )
            except Exception:
                pass


def wait_until_ready(target_id, timeout=3):
    """等待页面加载完成"""
    start = time.time()
    while time.time() - start < timeout:
        info = cdp_info(target_id)
        if info.get('ready') == 'complete':
            return True
        time.sleep(0.3)
    return False


def clean_text(text):
    """通用正文清洗"""
    # 页脚版权噪音（优先清除）
    text = re.sub(r'网站地图\s*版权保护\s*免责声明\s*关于我们', '', text)
    text = re.sub(r'版权所有[^\n]*', '', text)
    text = re.sub(r'地址[^\n]*', '', text)
    text = re.sub(r'邮编[^\n]*', '', text)
    text = re.sub(r'ICP备案[^\s]+', '', text)
    text = re.sub(r'网站标识码[^\s]*', '', text)
    text = re.sub(r'京公网安备[^\s]+', '', text)
    text = re.sub(r'电话[^\n]*', '', text)
    text = re.sub(r'分享到[^\n]*', '', text)
    text = re.sub(r'【返回顶部】【打印本页】【关闭窗口】', '', text)
    text = re.sub(r'上[^\n]*?篇：[^\n]*', '', text)
    text = re.sub(r'下[^\n]*?篇：[^\n]*', '', text)
    text = re.sub(r'按回车键[^\n]*导盲模式。', '', text)
    text = re.sub(r'编辑：[^\n]*', '', text)
    text = re.sub(r'时间：[^\n]*', '', text)
    text = re.sub(r'发布日期：[^\n]*', '', text)
    # 来源行可能有用，改为只删"来源："前缀，不删整行
    text = re.sub(r'^来源：.*$', '', text, flags=re.MULTILINE)
    text = re.sub(r'^当前位置[^\n]*', '', text, flags=re.MULTILINE)
    # 清理多余空白
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def extract_links(target_id, link_selectors):
    """从搜索结果页提取链接，遍历所有选择器并合并去重"""
    selectors = [s.strip() for s in link_selectors.split(',')]
    all_links = []
    seen = set()

    for sel in selectors:
        escaped_sel = sel.replace("'", "\\'")
        script = f"""var r=[];document.querySelectorAll('{escaped_sel}').forEach(function(a){{if(a.href)r.push({{text:a.textContent.replace(/<[^>]+>/g,'').trim(),href:a.href}})}});JSON.stringify(r)"""
        result = cdp_eval(target_id, script)
        if result:
            try:
                links = json.loads(result)
                for link in links:
                    # href去重（同选择器内可能重复）
                    if link['href'] not in seen:
                        seen.add(link['href'])
                        all_links.append(link)
            except Exception as e:
                print(f"    ⚠️ 选择器 '{sel}' 解析失败: {e}", flush=True)
    return all_links


def extract_article_text_fast(target_id, body_selectors_raw):
    """一次eval完成所有选择器尝试，返回命中结果"""
    selectors_list = [s.strip() for s in body_selectors_raw.split('|') if s.strip()]
    script = f"""
    (function(){{
        var selectors = {json.dumps(selectors_list)};
        for (var i = 0; i < selectors.length; i++) {{
            var sel = selectors[i];
            var el = document.querySelector(sel);
            if (!el) continue;
            var t = el.innerText.trim();
            if (t.length < 100) continue;
            var html = el.innerHTML;
            var hasTable = html.indexOf('<table') !== -1;
            var tables = [];
            if (hasTable) {{
                var tbls = el.querySelectorAll('table');
                tbls.forEach(function(tbl) {{
                    var rows = tbl.querySelectorAll('tr');
                    var data = [];
                    rows.forEach(function(row) {{
                        var cells = row.querySelectorAll('td, th');
                        var rowData = [];
                        cells.forEach(function(c) {{
                            rowData.push(c.textContent.replace(/\\s+/g, ' ').trim());
                        }});
                        if (rowData.length > 0) data.push(rowData);
                    }});
                    if (data.length > 0) tables.push(data);
                }});
            }}
            return JSON.stringify({{selector: sel, text: t, tables: tables}});
        }}
        var body = document.body.innerText.trim();
        return JSON.stringify({{selector: 'body(fallback)', text: body, tables: []}});
    }})()
    """
    result = cdp_eval(target_id, script)
    if result:
        try:
            data = json.loads(result)
            if data and data.get('text'):
                data['text'] = clean_text(data['text'])
                return data
        except Exception as e:
            print(f"    ⚠️ 正文解析失败: {e}", flush=True)
    return {'selector': 'body(fallback)', 'text': '', 'tables': []}


def generate_summary(title, text, max_chars=3000):
    """生成文章摘要 - 取中间段落避免噪音"""
    content = text[:max_chars] if len(text) > max_chars else text
    lines = [l.strip() for l in content.split('\n') if len(l.strip()) > 15]

    if not lines:
        return "（摘要提取失败，正文过短）"

    # 过滤短句（导航栏噪音）
    content_lines = [l for l in lines if len(l) >= 30]
    if len(content_lines) < 3:
        content_lines = lines

    total = len(content_lines)
    if total <= 3:
        key_lines = content_lines
    else:
        start = int(total * 0.4)
        end = int(total * 0.75)
        key_lines = content_lines[start:end]

    if not key_lines:
        return content_lines[0][:120] + '...' if len(content_lines[0]) > 120 else content_lines[0]

    selected = key_lines[:3]
    summary = ' '.join(selected)[:100]
    if len(' '.join(selected)) > 100:
        summary += '...'
    return summary


def crawl_and_display(urls, body_selectors, concurrency=5, max_retries=2):
    """
    流式爬取：每篇完成立即显示，支持摘要生成
    确保异常时也能清理已打开的tab
    失败自动重试（最多max_retries次）
    """
    total = len(urls)
    failed_urls = []  # 记录失败URL便于调试

    # 重试计数：{url: retry_count}
    retry_count = {}
    for url in urls:
        retry_count[url] = 0

    pending = list(urls)
    processed_urls = set()  # 记录本轮已处理完的URL（不含重试入队的）

    try:
        while pending:
            batch = pending[:concurrency]
            batch_num = (total - len(pending)) // concurrency + 1
            total_batches = (total + concurrency - 1) // concurrency

            print(f"\n📦 批次 {batch_num}/{total_batches} 正在处理...", flush=True)

            # 批量创建tab
            targets = [cdp_new(url) for url in batch]

            # 等待所有tab加载
            for t in targets:
                if t:
                    wait_until_ready(t, timeout=MAX_WAIT_READY)

            # 重置本轮已处理记录
            processed_urls.clear()

            # 逐个提取并实时显示
            for i, (url, t) in enumerate(zip(batch, targets)):
                idx = total - len(pending) + i + 1

                if not t:
                    rc = retry_count.get(url, 0)
                    if rc < max_retries:
                        print(f"    🔄 [{idx}/{total}] 重试({rc+1}/{max_retries}) - 无法打开tab", flush=True)
                        pending.append(url)
                        retry_count[url] = rc + 1
                    else:
                        failed_urls.append((url, '无法创建tab'))
                        print(f"    ⏭️  [{idx}/{total}] 跳过（无法打开）", flush=True)
                        processed_urls.add(url)
                    continue

                try:
                    data = extract_article_text_fast(t, body_selectors)
                    text = data.get('text', '') if data else ''
                    tables = data.get('tables', []) if data else []
                    selector_used = data.get('selector', 'unknown') if data else 'unknown'
                finally:
                    # 必须关闭tab
                    cdp_close(t)

                if text and len(text) > 50:
                    summary = generate_summary(url, text)
                    result = {
                        'url': url,
                        'text': text,
                        'tables': tables,
                        'selector': selector_used,
                        'summary': summary,
                        'char_count': len(text)
                    }
                    print(f"    ✅ [{len(processed_urls)+1}/{total}] {selector_used[:15]:<15} {len(text):>5}字", flush=True)
                    processed_urls.add(url)
                    yield result
                else:
                    rc = retry_count.get(url, 0)
                    if rc < max_retries:
                        print(f"    🔄 [{idx}/{total}] 重试({rc+1}/{max_retries}) - 内容太短", flush=True)
                        pending.append(url)
                        retry_count[url] = rc + 1
                    else:
                        failed_urls.append((url, '内容为空或太短'))
                        print(f"    ❌ [{idx}/{total}] 内容为空或太短（已重试{max_retries}次）", flush=True)
                        processed_urls.add(url)
                        yield None

            # 只移除真正处理完的URL（排除重试入队的）
            for url in batch:
                if url in processed_urls and url in pending:
                    pending.remove(url)

    finally:
        # 确保批次结束后清理所有tab
        pass

    # 报告统计
    success = total - len(failed_urls)
    print(f"\n🎉 抓取完成！成功 {success}/{total} 篇", flush=True)
    if failed_urls:
        print(f"   失败 {len(failed_urls)} 条:", flush=True)
        for url, reason in failed_urls[:3]:
            print(f"     - {url[:50]}... ({reason})", flush=True)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
        is_external=True
    )
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def generate_doc(links, results, output_path, keyword, search_url):
    doc = Document()
    domain = re.sub(r'https?://', '', search_url.split('/')[0] if '://' in search_url else search_url)
    title = doc.add_heading(f'"{keyword}" 搜索结果汇编', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'共收录 {len(results)} 篇文章\n（来源：{domain} 搜索"{keyword}"第1页）')
    doc.add_paragraph()
    count = 0
    for l in links:
        url = l['href']
        title_text = l['text']
        data = results.get(url, {})
        if not data:
            continue
        text = data.get('text', '')
        tables = data.get('tables', [])
        if not text or len(text) < 50:
            continue
        count += 1
        doc.add_heading(f'{count}. {title_text}', level=2)
        p = doc.add_paragraph()
        p.add_run('原文链接：').bold = True
        add_hyperlink(p, url, url)
        if tables:
            doc.add_heading('（正文）', level=3)
            doc.add_paragraph(f'（含 {len(tables)} 个表格）')
            for ti, table_data in enumerate(tables):
                if not table_data:
                    continue
                rows = len(table_data)
                cols = max(len(row) for row in table_data) if table_data else 0
                if rows > 0 and cols > 0:
                    doc.add_paragraph(f'【表格 {ti+1}】')
                    word_table = doc.add_table(rows=rows, cols=cols)
                    word_table.style = 'Table Grid'
                    for ri, row_data in enumerate(table_data):
                        for ci, cell_text in enumerate(row_data):
                            if ci < cols:
                                cell = word_table.cell(ri, ci)
                                cell.text = cell_text
                    doc.add_paragraph()
        else:
            doc.add_heading('（正文）', level=3)
            if len(text) > 10000:
                text = text[:10000] + '...[已截断]'
            doc.add_paragraph(text)
            doc.add_paragraph()
    doc.save(output_path)
    print(f"\n文档已保存: {output_path}")
    print(f"共 {count} 篇文章")


def stream_analyze_output(search_url_template, keyword, link_selectors=None, body_selectors=None):
    """流式分析模式：边抓边显示，每篇带摘要"""
    global _open_tabs

    if link_selectors is None:
        link_selectors = DEFAULT_LINK_SELECTORS
    if body_selectors is None:
        body_selectors = DEFAULT_BODY_SELECTORS

    keyword_encoded = urllib.parse.quote(keyword)
    search_url = search_url_template.replace('{keyword}', keyword_encoded).replace('{kw}', keyword_encoded)

    print("=" * 60, flush=True)
    print(f"🔍 搜索引擎: {search_url_template}", flush=True)
    print(f"📌 关键词: {keyword}", flush=True)
    print(f"📊 模式: 流式抓取 + 摘要生成", flush=True)
    print("=" * 60, flush=True)

    try:
        # Step 1: 打开搜索页
        print(f"\n🌐 [1/3] 打开搜索页...", flush=True)
        search_target = cdp_new(search_url)
        if not search_target:
            print("❌ ERROR: 无法打开搜索页，请确认Chrome CDP正在运行（http://localhost:3456）")
            return

        if not wait_until_ready(search_target, timeout=MAX_WAIT_READY):
            print("❌ 搜索页加载超时", flush=True)
            return

        # Step 2: 提取链接
        print(f"📋 [2/3] 提取搜索结果...", flush=True)
        links = extract_links(search_target, link_selectors)
        cdp_close(search_target)  # 搜索页用完即关闭

        if not links:
            print("❌ 未找到任何链接，请检查链接选择器是否正确")
            return

        # 去重
        original_count = len(links)
        links = dedup_links(links)
        dedup_count = original_count - len(links)
        print(f"   ✅ 找到 {original_count} 条结果", flush=True)
        if dedup_count > 0:
            print(f"   🔄 去重后 {len(links)} 条（去除重复 {dedup_count} 条）", flush=True)

        # 限制处理数量
        max_articles = getattr(sys, 'max_articles', 0) or len(links)
        if max_articles > 0 and max_articles < len(links):
            links = links[:max_articles]
            print(f"   📌 限制处理前 {max_articles} 条", flush=True)

        # 打印链接预览
        print(f"\n   链接预览（前10条）:", flush=True)
        for i, l in enumerate(links[:10]):
            norm_tag = f" [→{l.get('_norm_url','')[:30]}...]" if '_norm_url' in l else ''
            print(f"   {i+1:2d}. {l['text'][:40]}{'...' if len(l['text']) > 40 else ''}{norm_tag}", flush=True)
        if len(links) > 10:
            print(f"   ... 还有 {len(links)-10} 条\n", flush=True)

        # Step 3: 流式抓取
        print(f"📥 [3/3] 开始流式抓取（{len(links)} 篇，{CONCURRENCY}并发）...", flush=True)
        print("-" * 60, flush=True)

        urls = [l['href'] for l in links]
        all_results = {}

        for result in crawl_and_display(urls, body_selectors, concurrency=CONCURRENCY, max_retries=DEFAULT_MAX_RETRIES):
            if result:
                all_results[result['url']] = {
                    'text': result['text'],
                    'tables': result.get('tables', []),
                    'selector': result.get('selector', ''),
                    'summary': result.get('summary', ''),
                    'char_count': result.get('char_count', 0)
                }

        print("-" * 60, flush=True)
        print(f"\n📊 抓取完成！成功 {len(all_results)}/{len(urls)} 篇\n", flush=True)

        # 输出汇总（供AI分析）
        print("=" * 60, flush=True)
        print("📑 内容摘要汇总（供AI分析）", flush=True)
        print("=" * 60, flush=True)

        for i, l in enumerate(links):
            url = l['href']
            title_text = l['text']
            data = all_results.get(url)
            if not data:
                continue

            summary = data.get('summary', '')
            text = data.get('text', '')
            norm_url = l.get('_norm_url', '')

            print(f"\n{'='*60}")
            print(f"【{i+1}】{title_text}")
            if norm_url:
                print(f"🔗 真实URL: {norm_url[:70]}{'...' if len(norm_url) > 70 else ''}")
            print(f"📝 摘要: {summary}")
            print(f"{'='*60}")
            # 截取正文前1500字用于分析
            display_text = text[:1500] + "..." if len(text) > 1500 else text
            print(display_text)
            print()

    finally:
        # 确保异常时也清理所有tab
        cleanup_all_tabs()


def main(search_url_template, keyword, output_dir=None, link_selectors=None, body_selectors=None, mode='analyze'):
    global _open_tabs

    if mode == 'analyze':
        return stream_analyze_output(search_url_template, keyword, link_selectors, body_selectors)

    # doc模式
    if link_selectors is None:
        link_selectors = DEFAULT_LINK_SELECTORS
    if body_selectors is None:
        body_selectors = DEFAULT_BODY_SELECTORS

    keyword_encoded = urllib.parse.quote(keyword)
    search_url = search_url_template.replace('{keyword}', keyword_encoded).replace('{kw}', keyword_encoded)

    print("=" * 50)
    print(f"搜索URL: {search_url_template}")
    print(f"关键词: {keyword}")
    print(f"模式: 生成Word文档")
    print("=" * 50)

    try:
        print("\n[1/4] 打开搜索页...", flush=True)
        search_target = cdp_new(search_url)
        if not search_target:
            print("ERROR: 无法打开搜索页，请确认Chrome CDP正在运行（http://localhost:3456）")
            sys.exit(1)

        wait_until_ready(search_target, timeout=MAX_WAIT_READY)

        print("\n[2/4] 提取第1页搜索结果...", flush=True)
        links = extract_links(search_target, link_selectors)
        cdp_close(search_target)

        if not links:
            print("  未找到任何链接，请检查链接选择器是否正确")
            sys.exit(1)

        # 去重
        links = dedup_links(links)
        print(f"  共提取到 {len(links)} 条链接（去重后）")

        # 限制处理数量
        max_articles = getattr(sys, 'max_articles', 0) or len(links)
        if max_articles > 0 and max_articles < len(links):
            links = links[:max_articles]
            print(f"  限制处理前 {max_articles} 条")

        for i, l in enumerate(links):
            print(f"  {i+1:3d}. {l['text'][:50]}")

        print(f"\n[3/4] 批量爬取正文 ({len(links)} 篇，{CONCURRENCY}并发)...", flush=True)
        urls = [l['href'] for l in links]

        results = {}
        for result in crawl_and_display(urls, body_selectors, concurrency=CONCURRENCY, max_retries=DEFAULT_MAX_RETRIES):
            if result:
                results[result['url']] = {
                    'text': result['text'],
                    'tables': result.get('tables', []),
                    'selector': result.get('selector', '')
                }

        print(f"  成功爬取 {len(results)} / {len(urls)} 篇")

        if output_dir is None:
            safe_keyword = re.sub(r'[\\/:*?"<>|]', '_', keyword)
            output_dir = os.path.expanduser(f"~/Desktop/{safe_keyword}搜索结果")
        os.makedirs(output_dir, exist_ok=True)
        safe_keyword = re.sub(r'[\\/:*?"<>|]', '_', keyword)
        output_path = os.path.join(output_dir, f'"{keyword}"搜索结果.docx')
        print(f"\n[4/4] 生成Word文档...", flush=True)
        generate_doc(links, results, output_path, keyword, search_url_template)
        print("完成！")

    finally:
        cleanup_all_tabs()


if __name__ == '__main__':
    # 注册信号处理器，确保Ctrl+C时也能清理tab
    import signal

    def signal_handler(sig, frame):
        cleanup_all_tabs()
        sys.exit(0)

    signal.signal(signal.SIGINT, signal_handler)
    signal.signal(signal.SIGTERM, signal_handler)

    if len(sys.argv) < 3:
        print("用法: python3 general_search.py <搜索URL> <关键词> [输出目录] [链接选择器] [正文选择器] [analyze]")
        print("")
        print("参数说明:")
        print("  搜索URL     : 搜索页URL，{keyword} 或 {kw} 占位符会被关键词替换")
        print("  关键词       : 要搜索的关键词")
        print("  输出目录     :（可选）文档输出路径，默认 ~/Desktop/<关键词>搜索结果")
        print("  链接选择器  :（可选）CSS选择器提取搜索结果链接，支持逗号分隔多选择器")
        print("  正文选择器  :（可选）|分隔的CSS选择器列表")
        print("  analyze     :（可选）流式输出+摘要模式，AI分析专用")
        print("")
        print("示例:")
        print('  python3 general_search.py "https://www.baidu.com/s?wd={keyword}" "人工智能"')
        print('  python3 general_search.py "https://www.baidu.com/s?wd={keyword}" "人工智能" ~/Desktop analyze')
        print('  python3 general_search.py "https://www.baidu.com/s?wd={keyword}" "新闻" "" ".news a, .article a" "article|p" analyze')
        sys.exit(1)

    search_url = sys.argv[1]
    keyword = sys.argv[2]
    output_dir = sys.argv[3] if len(sys.argv) > 3 else None
    link_selectors = sys.argv[4] if len(sys.argv) > 4 else None
    body_selectors = sys.argv[5] if len(sys.argv) > 5 else None
    mode = sys.argv[6] if len(sys.argv) > 6 else 'analyze'
    sys.max_articles = int(sys.argv[7]) if len(sys.argv) > 7 and sys.argv[7].isdigit() else 0

    main(search_url, keyword, output_dir, link_selectors, body_selectors, mode)
